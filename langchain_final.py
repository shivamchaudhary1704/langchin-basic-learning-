import pandas as pd
from pathlib import Path
import re
import datetime

from langchain_community.llms import Ollama
from langchain_community.tools import DuckDuckGoSearchRun, WikipediaQueryRun
from langchain_community.utilities import DuckDuckGoSearchAPIWrapper, WikipediaAPIWrapper
from langchain.agents import initialize_agent, Tool, AgentType
from docx import Document
from docx.shared import Pt
llm = Ollama(model="llama3", temperature=0.2)

ddg = DuckDuckGoSearchAPIWrapper()
ddg_tool = DuckDuckGoSearchRun(api_wrapper=ddg)

wiki = WikipediaAPIWrapper(top_k_results=5)
wiki_tool = WikipediaQueryRun(api_wrapper=wiki)

tools = [
    Tool(
        name="DuckDuckGo",
        func=ddg_tool.run,
        description="Use for general web search and recent info"
    ),
    Tool(
        name="Wikipedia",
        func=wiki_tool.run,
        description="Use for background info, definitions, and summaries"
    )
]
agent = initialize_agent(
    tools=tools,
    llm=llm,
    agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
    verbose=True
)
def generate_report(topic: str, out_path: Path = None) -> Path:
    system_prompt = f"""
You are a research assistant. Create a structured research report on: "{topic}"
- Use headings: Introduction, Background, Current Trends, Key Data/Stats, Opportunities/Risks, Conclusion
- Cite sources inline if possible
- Make it concise but informative
- Use info from Wikipedia and web search (DuckDuckGo)
"""

    # Run agent
    report_text = agent.run(system_prompt)

    # -----------------------------
    # Convert to DOCX
    # -----------------------------
    out_name = out_path or f"{re.sub(r'[^A-Za-z0-9]+','_',topic)[:50]}_report.docx"
    doc = Document()
    doc.add_heading(f"Research Report: {topic}", 0)
    for line in report_text.splitlines():
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", "").strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", "").strip(), level=1)
        else:
            doc.add_paragraph(line.strip())
    doc.add_paragraph("")
    doc.add_paragraph(f"Generated locally with Llama 3 · {datetime.date.today().isoformat()}")
    doc.save(out_name)

    print(f"\n✅ Report saved as: {out_name}")
    return Path(out_name) 

# -----------------------------
# 5️⃣ Example Usage
# -----------------------------
if __name__ == "__main__":
    topic_input="weather in noida"
    # topic_input = input("Enter your research topic: ")
    generate_report(topic_input)

