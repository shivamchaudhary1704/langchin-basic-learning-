from pathlib import Path
import re
import datetime
import sys
import argparse

# --- Llama (Ollama) via LangChain community ---
from langchain_community.llms import Ollama

# --- DOCX ---
from docx import Document

# -----------------------------
# Config
# -----------------------------
MODEL_NAME = "llama3"   # change to your local model tag if different
TEMPERATURE = 0.2

llm = Ollama(model=MODEL_NAME, temperature=TEMPERATURE)

# -----------------------------
# Prompt Builder
# -----------------------------
def build_prompt(topic: str) -> str:
    return f"""
You are a research assistant. Create a structured research report on: "{topic}"

Requirements:
- Sections (use these exact headings): Introduction, Background, Current Trends, Key Data/Stats, Opportunities/Risks, Conclusion
- Be concise but informative (400–800 words total)
- Use bullet points where helpful
- If you mention facts, include brief inline source hints like (source: org / paper / year) — no web calls needed
- Output in Markdown with '#' and '##' headings only
"""

# -----------------------------
# Llama-only generation
# -----------------------------
def generate_text(topic: str) -> str:
    prompt = build_prompt(topic)
    # LangChain's Ollama: use .invoke (preferred) or .predict
    try:
        return llm.invoke(prompt)
    except AttributeError:
        return llm.predict(prompt)

# -----------------------------
# Save DOCX
# -----------------------------
def markdown_to_docx(md_text: str, title: str, out_path: Path | None = None) -> Path:
    # safe filename
    default_name = f"{re.sub(r'[^A-Za-z0-9]+','_',title)[:50]}_report.docx"
    out_path = Path(out_path) if out_path else Path(default_name)

    doc = Document()
    doc.add_heading(f"Research Report: {title}", 0)

    for raw_line in md_text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        else:
            # handle bullets
            if line.startswith("- "):
                p = doc.add_paragraph(line[2:].strip(), style=None)
                # Simple bullet by prefix; Word will render as normal paragraph unless styled
                # If you want actual bullet style:
                p.style = "List Bullet"
            else:
                doc.add_paragraph(line)

    doc.add_paragraph("")
    doc.add_paragraph(f"Generated locally with Llama · {datetime.date.today().isoformat()}")
    doc.save(out_path)
    print(f"✅ Report saved as: {out_path}")
    return out_path

# -----------------------------
# Public API
# -----------------------------
def generate_report(topic: str, out_path: Path | None = None) -> Path:
    text = generate_text(topic)
    return markdown_to_docx(text, topic, out_path)

# -----------------------------
# CLI
# -----------------------------
def main():
    parser = argparse.ArgumentParser(description="Generate a research report using local Llama (Ollama).")
    parser.add_argument("--topic", type=str, help="Research topic")
    parser.add_argument("--out", type=str, help="Output .docx path")
    args = parser.parse_args()

    if args.topic:
        topic = args.topic
    else:
        # Interactive fallback (avoids VS Code read-only terminal issue by allowing CLI usage)
        try:
            topic = input("Enter your research topic: ").strip()
        except Exception:
            print("No interactive input available. Provide --topic \"Your Topic\".", file=sys.stderr)
            sys.exit(1)

    generate_report(topic, Path(args.out) if args.out else None)

if __name__ == "__main__":
    main()
